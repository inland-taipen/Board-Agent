"""DOCX ingestion via python-docx.

Word has no page model that a parser can read - pagination is a rendering
decision made by Word itself. Rather than invent page numbers that would not
match the printed board pack, we count *explicit* page breaks and cite
paragraph ranges alongside an approximate page. A director following a
citation sees "p. ~3 (paras 41-58)", which is honest about its precision.
"""

from __future__ import annotations

from pathlib import Path

import docx
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .base import Segment, normalise

# Target size of a pseudo-page when the document has no explicit page breaks.
_CHARS_PER_PAGE = 2800


def parse(path: Path) -> list[Segment]:
    document = docx.Document(str(path))
    segments: list[Segment] = []

    page = 1
    para_index = 0
    buf: list[str] = []
    buf_start = 1
    heading: str | None = None

    def flush() -> None:
        nonlocal buf, buf_start, heading
        if not buf:
            return
        text = normalise("\n".join(buf))
        if text:
            segments.append(
                Segment(
                    page=page,
                    text=text,
                    locator=f"p. ~{page} (paras {buf_start}-{para_index})",
                    heading=heading,
                    kind="body",
                )
            )
        buf = []
        buf_start = para_index + 1

    for block in _iter_blocks(document):
        if isinstance(block, Table):
            flush()
            table_text = _render_table(block)
            if table_text:
                segments.append(
                    Segment(
                        page=page,
                        text=table_text,
                        locator=f"p. ~{page}, table after para {para_index}",
                        heading=heading,
                        kind="table",
                    )
                )
            continue

        para_index += 1
        style = (block.style.name or "") if block.style is not None else ""
        text = block.text.strip()

        if _has_page_break(block):
            flush()
            page += 1
            buf_start = para_index

        if style.startswith("Heading") or style == "Title":
            flush()
            if text:
                heading = text
                buf.append(text)
            continue

        if text:
            buf.append(text)

        # No explicit breaks in this document - approximate by volume so that
        # locators still cluster the way a reader would expect.
        if sum(len(x) for x in buf) > _CHARS_PER_PAGE:
            flush()
            page += 1

    flush()
    return segments


def _iter_blocks(document: DocxDocument):
    """Yield paragraphs and tables in true document order.

    python-docx exposes `.paragraphs` and `.tables` as separate sequences,
    which loses the interleaving - a risk register table would otherwise be
    detached from the paragraph that introduces it.
    """
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _has_page_break(paragraph: Paragraph) -> bool:
    for br in paragraph._p.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    # Word records where it last rendered a page break; when present it is the
    # closest thing to real pagination a parser can see.
    return any(True for _ in paragraph._p.iter(qn("w:lastRenderedPageBreak")))


def _render_table(table: Table) -> str:
    lines: list[str] = []
    for row in table.rows:
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        if any(cells):
            lines.append(" | ".join(cells))
    return "\n".join(lines)
