"""DOCX rendering.

Company secretaries edit the briefing before circulating it, so the DOCX is the
working copy: real Word styles, a proper heading outline that generates a
navigation pane, and a classification banner in the header of every page.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .layout import BriefingDocument, build

_INK = RGBColor(0x1A, 0x1A, 0x1A)
_MUTED = RGBColor(0x5A, 0x5A, 0x5A)
_ACCENT = RGBColor(0x0B, 0x3D, 0x5C)
_ALERT = RGBColor(0x8B, 0x1A, 0x1A)


def render(doc: BriefingDocument, destination: Path) -> Path:
    document = Document()
    _page_setup(document)
    _styles(document)
    _header_footer(document, doc)

    for block in build(doc):
        _render_block(document, block)

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))
    return destination


def _page_setup(document: Document) -> None:
    section = document.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)


def _styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = _INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15


def _header_footer(document: Document, doc: BriefingDocument) -> None:
    section = document.sections[0]

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(doc.classification_banner)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = _ALERT

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"{doc.company}  |  {doc.meeting_label}  |  Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = _MUTED
    _page_number_field(footer)


def _page_number_field(paragraph) -> None:
    """Insert a live PAGE field so pagination survives editing in Word."""
    run = paragraph.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = _MUTED

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _render_block(document: Document, block) -> None:
    kind = block.kind

    if kind == "pagebreak":
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        return

    if kind == "title":
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(block.text)
        run.bold = True
        run.font.size = Pt(26)
        run.font.color.rgb = _ACCENT
        return

    if kind == "subtitle":
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(block.text)
        run.font.size = Pt(14)
        run.font.color.rgb = _MUTED
        return

    if kind == "h1":
        heading = document.add_heading(block.text, level=1)
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(8)
        for run in heading.runs:
            run.font.color.rgb = _ACCENT
            run.font.size = Pt(15)
        _rule(document)
        return

    if kind == "h2":
        heading = document.add_heading(block.text, level=2)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(4)
        for run in heading.runs:
            run.font.color.rgb = _INK
            run.font.size = Pt(12)
        return

    if kind == "para":
        document.add_paragraph(block.text)
        return

    if kind == "kv":
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        label = p.add_run(f"{block.label}:  ")
        label.bold = True
        label.font.color.rgb = _ACCENT
        p.add_run(block.text or "not stated")
        return

    if kind == "note":
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(block.text)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = _MUTED
        return

    if kind == "bullets":
        for item in block.items or []:
            document.add_paragraph(item, style="List Bullet")
        return

    if kind == "cite":
        items = block.items or []
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.left_indent = Inches(0.2)
        run = p.add_run("Source: " + ("; ".join(items) if items else "no source cited"))
        run.italic = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = _MUTED if items else _ALERT
        return

    if kind == "table":
        _render_table(document, block)
        return


def _render_table(document: Document, block) -> None:
    rows = block.items or []
    if not rows:
        return

    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    for cell, value in zip(table.rows[0].cells, rows[0], strict=False):
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(value))
        run.bold = True
        run.font.size = Pt(9)

    for row_values in rows[1:]:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values, strict=False):
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(value))
            run.font.size = Pt(8.5)

    document.add_paragraph()


def _rule(document: Document) -> None:
    """Thin horizontal rule under section headings."""
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0B3D5C")
    borders.append(bottom)
    pPr.append(borders)
