"""DOCX and PDF export tests.

The exports are what actually reaches directors, so the assertions check the
things a reader would notice: the classification stamp, the mandated five
sections, the citations, and the fact that an unresolved reference is shown as
unresolved rather than quietly dropped.
"""

from __future__ import annotations

from fixtures import BRIEFING

from boardlens.brief.verify import verify
from boardlens.export import BriefingDocument, render_docx, render_pdf
from boardlens.export.layout import build, citations_for
from boardlens.ingest import classify, parse_file
from boardlens.rag import PackIndex, chunk_segments


def _document(sample_pack) -> BriefingDocument:
    chunks, counter = [], 0
    for path in sorted(sample_pack.iterdir()):
        segments = parse_file(path)
        kind = classify(path.name, "\n".join(s.text for s in segments[:6]))
        produced = chunk_segments(
            segments,
            doc_id=path.stem,
            doc_name=path.name,
            doc_kind=str(kind),
            start_index=counter,
        )
        counter += len(produced)
        chunks.extend(produced)

    index = PackIndex(chunks).build()
    report = verify(BRIEFING, index)

    return BriefingDocument(
        company="Meridian Industries Limited",
        meeting_label="119th Board Meeting",
        meeting_date="2026-08-21",
        classification="strictly_confidential",
        generated_at="2026-08-20",
        model="claude-opus-5",
        briefing=BRIEFING.model_dump(),
        citation_map=report.citation_map,
        verification=report.to_dict(),
        action_register=[
            {
                "action": "Standardise the ESG disclosure framework across all subsidiaries",
                "owner": "Chief Sustainability Officer",
                "raised_at": "116th meeting",
                "committed_date": "30 June 2026",
                "status": "unclear",
                "ageing_cycles": 2,
            }
        ],
    )


def test_layout_covers_the_five_mandated_sections(sample_pack):
    blocks = list(build(_document(sample_pack)))
    headings = [b.text for b in blocks if b.kind == "h1"]

    assert "1. Critical risks for board attention" in headings
    assert "2. Unresolved actions from previous meetings" in headings
    assert "3. Material performance changes" in headings
    assert "4. Questions the board should put to management" in headings
    assert "5. Decisions required at this meeting" in headings
    assert "Coverage and limitations" in headings
    assert "Annexure A - Complete action register" in headings
    assert "Annexure B - Source index" in headings


def test_unresolved_reference_is_shown_not_hidden():
    rendered = citations_for(["c0001", "c9999"], {"c0001": {"document": "Deck.pptx", "locator": "Slide 3"}})

    assert rendered == ["Deck.pptx, Slide 3", "[unverified reference c9999]"]


def test_citations_deduplicate_but_keep_order():
    cmap = {
        "c1": {"document": "MIS.xlsx", "locator": "Sheet 'P&L', rows 1-13"},
        "c2": {"document": "MIS.xlsx", "locator": "Sheet 'P&L', rows 1-13"},
        "c3": {"document": "Deck.pptx", "locator": "Slide 3"},
    }
    assert citations_for(["c1", "c2", "c3"], cmap) == [
        "MIS.xlsx, Sheet 'P&L', rows 1-13",
        "Deck.pptx, Slide 3",
    ]


def test_docx_export_is_readable_and_stamped(sample_pack, tmp_path):
    import docx

    path = render_docx(_document(sample_pack), tmp_path / "briefing.docx")
    assert path.exists() and path.stat().st_size > 10_000

    document = docx.Document(str(path))
    text = "\n".join(p.text for p in document.paragraphs)

    assert "Board Briefing" in text
    assert "Meridian Industries Limited" in text
    assert "1. Critical risks for board attention" in text
    assert "Covenant headroom on debt service coverage" in text
    assert "Source:" in text
    assert "[unverified reference c9999]" in text

    header = document.sections[0].header.paragraphs[0].text
    assert "STRICTLY CONFIDENTIAL" in header

    # The source index annexure is a table, not paragraphs.
    assert document.tables, "no annexure tables rendered"


def test_pdf_export_paginates_and_embeds_metadata(sample_pack, tmp_path):
    from pypdf import PdfReader

    path = render_pdf(_document(sample_pack), tmp_path / "briefing.pdf")
    assert path.exists() and path.stat().st_size > 5_000

    reader = PdfReader(str(path))
    assert len(reader.pages) >= 3

    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    assert "Board Briefing" in text
    assert "STRICTLY CONFIDENTIAL" in text
    assert "Covenant headroom" in text
    assert "Annexure B" in text


def test_export_escapes_ampersands_and_angle_brackets(sample_pack, tmp_path):
    """Board text contains '&' constantly; unescaped it aborts the PDF render."""
    document = _document(sample_pack)
    document.briefing["critical_risks"][0]["title"] = "Ind AS 116 <lease> & covenant risk"

    path = render_pdf(document, tmp_path / "escaped.pdf")
    assert path.exists() and path.stat().st_size > 5_000
